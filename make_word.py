from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_word():
    doc = Document()

    # Title
    title = doc.add_heading('Contoh Dokumen Word dari Cuy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    p = doc.add_paragraph()
    p.add_run('Halo Callunk! 👑').bold = True
    p.add_run('\nIni adalah contoh file Word yang dibuat secara otomatis oleh aku menggunakan Python di VPS-mu.')

    # List of capabilities
    doc.add_heading('Apa saja yang bisa aku lakukan sekarang?', level=1)
    capabilities = [
        'Membuat dokumen Word (.docx) secara otomatis.',
        'Menyusun laporan formal dengan struktur yang rapi.',
        'Membuat tabel data di dalam dokumen.',
        'Mengatur format teks (Bold, Italic, Alignment).',
        'Mengotomatisasi pembuatan surat atau SK.'
    ]
    for item in capabilities:
        doc.add_paragraph(item, style='List Bullet')

    # Closing
    doc.add_paragraph('\nSekarang aku sudah punya "tangan" untuk membuat file Word. Tidak perlu copy-paste manual lagi! 🚀')

    # Save
    doc.save('Contoh_Dokumen_Cuy.docx')
    print("File created successfully!")

if __name__ == "__main__":
    create_sample_word()
