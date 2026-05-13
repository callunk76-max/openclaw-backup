#!/usr/bin/env python3
"""Convert LPSE report PDF → DOCX with Kabupaten Bulukumba attributes."""
import os
from pdfminer.high_level import extract_text
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PDF_PATH = "/root/.openclaw/workspace/laporan_kegiatan_lpse.pdf"
DOCX_PATH = "/root/.openclaw/workspace/Laporan_Kegiatan_LPSE_Bulukumba.docx"

text = extract_text(PDF_PATH)

# Replacements
repl = {
    "PEMERINTAH KOTA BANJARBARU": "PEMERINTAH KABUPATEN BULUKUMBA",
    "Pemerintah Kota Banjarbaru": "Pemerintah Kabupaten Bulukumba",
    "KOTA BANJARBARU": "KABUPATEN BULUKUMBA",
    "Kota Banjarbaru": "Kabupaten Bulukumba",
    "Banjarbaru": "Bulukumba",
    "BANJARBARU": "BULUKUMBA",
    "Bagian Pengadaan Barang dan Jasa": "LPSE Kabupaten Bulukumba",
    "BAGIAN PENGADAAN BARANG DAN JASA": "LPSE KABUPATEN BULUKUMBA",
    "Hj. Renyta Setyawati, SP., MT.": "Kepala LPSE",
    "Pembina Tingkat I\nNIP. 19711010 199803 2 009": "",
    "NIP. 19711010 199803 2 009": "",
}
for old, new in repl.items():
    text = text.replace(old, new)

lines = text.split('\n')

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# === COVER PAGE ===
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PEMERINTAH KABUPATEN BULUKUMBA")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SEKRETARIAT DAERAH")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LPSE KABUPATEN BULUKUMBA")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LAPORAN KEGIATAN")
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LAYANAN PENGADAAN SECARA ELEKTRONIK")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("TRIWULAN I")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PERIODE BULAN JANUARI S/D MARET 2025")
run.font.size = Pt(12)

for _ in range(8):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PEMERINTAH KABUPATEN BULUKUMBA")
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2025")
run.bold = True
run.font.size = Pt(12)

doc.add_page_break()

# === CONTENT ===
body_text = []
katapengantar = False
for line in lines:
    stripped = line.strip()
    if not stripped:
        continue
    
    # Skip cover lines
    if stripped in ["PEMERINTAH KABUPATEN BULUKUMBA", "SEKRETARIAH DAERAH",
                    "LPSE KABUPATEN BULUKUMBA", "LAPORAN KEGIATAN",
                    "LAYANAN PENGADAAN", "SECARA ELEKTRONIK",
                    "TRIWULAN I", "PERIODE BULAN JANUARI S/D MARET 2025",
                    "2025", "i", "ii", "iii", "iv", "v", "vi", "vii"]:
        continue
    
    # Skip header/footer repeats
    if "LAPORAN KEGIATAN LAYANAN PENGADAAN SECARA ELEKTRONIK" in stripped and "TRIWULAN I" in stripped:
        continue
    if "PEMERINTAH KABUPATEN BULUKUMBA" in stripped and "2025" in stripped:
        continue
    if "LPSE KABUPATEN BULUKUMBA" in stripped and "UNIT KERJA" in stripped:
        continue
    if stripped.startswith("BAGIAN PENGADAAN") and "UNIT KERJA" in stripped:
        continue
    
    body_text.append(stripped)

# Write content
for line in body_text:
    # Headers
    if line in ["KATA PENGANTAR", "DAFTAR ISI", "BAB I.   PENDAHULUAN", 
                 "BAB I. PENDAHULUAN", "BAB II", "BAB II.", "BAB II. PENUTUP",
                 "PENUTUP"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        continue
    
    # Sub-headers (A., B., C., etc.)
    if line and len(line) < 80 and (line[0].isalpha() and line[1:3] in ['. ', '.  '] or line.startswith('b.') or line.startswith('c.') or line.startswith('A.') or line.startswith('B.') or line.startswith('C.') or line.startswith('D.')):
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        continue
    
    # Table of Contents page - skip dots
    if '...........................................' in line or '.......................................................' in line:
        continue
    
    # Normal paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run.space_after = Pt(6)

# === SIGNATURE SECTION ===
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Bulukumba,      April 2025")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Kepala LPSE Kabupaten Bulukumba,")
run.font.size = Pt(12)
run.font.name = 'Times New Roman'

for _ in range(3):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("_________________________")
run.bold = True
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("NIP. -")
run.font.size = Pt(11)

doc.save(DOCX_PATH)
print(f"✅ Saved: {DOCX_PATH}")
print(f"   Size: {os.path.getsize(DOCX_PATH) / 1024:.1f} KB")
