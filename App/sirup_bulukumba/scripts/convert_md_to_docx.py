
from docx import Document
from docx.shared import Pt
import os

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('- [ ] ') or line.startswith('- [x] ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"Converted {md_path} to {docx_path}")

files_to_convert = {
    "LPSE/Plan_2026/ROADMAP_STANDARISASI_2026.md": "LPSE/Plan_2026/ROADMAP_STANDARISASI_2026.docx",
    "LPSE/Plan_2026/CHECKLIST_KEDOKUMENAN.md": "LPSE/Plan_2026/CHECKLIST_KEDOKUMENAN.docx",
    "LPSE/Plan_2026/GUIDELINES_DOCUMENTATION.md": "LPSE/Plan_2026/GUIDELINES_DOCUMENTATION.docx"
}

base_path = "/root/.openclaw/workspace/github_backup/"
for md, docx in files_to_convert.items():
    md_to_docx(os.path.join(base_path, md), os.path.join(base_path, docx))
