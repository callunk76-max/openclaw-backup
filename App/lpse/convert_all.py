import os
import glob
from docx import Document
from docx.shared import Pt

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph() # Preserve empty lines as paragraph breaks
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- [ ] ') or line.startswith('- [x] ') or line.startswith('- '):
            clean_line = line.replace('- [ ] ', '[ ] ').replace('- [x] ', '[x] ').replace('- ', '')
            doc.add_paragraph(clean_line, style='List Bullet')
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print(f"✅ Converted: {os.path.basename(md_path)} -> {os.path.basename(docx_path)}")

base_dir = "/root/.openclaw/workspace/LPSE/Rencana_Standarisasi/AKSI"
md_files = glob.glob(f"{base_dir}/**/*.md", recursive=True)

count = 0
for md_file in md_files:
    docx_file = md_file.replace('.md', '.docx')
    md_to_docx(md_file, docx_file)
    count += 1

print(f"\nTotal {count} files converted successfully!")
